#!/usr/bin/env python3
# Copyright (c) 2025 WSO2 LLC.
# Generates a Word (.docx) visual operations guide with screenshot placeholders.
#
# Usage (from repo root or this directory):
#   python3 build_visual_guide_docx.py
#
# Output: Security-Advisory-Patches-Portal-Visual-Guide.docx (in this folder)
#
# Dependencies: pip install python-docx pillow

from __future__ import annotations

import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError as e:
    raise SystemExit("Install Pillow: pip install pillow") from e


def _load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        p = Path(path)
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def make_placeholder_image(out_path: Path, title: str, hint: str) -> None:
    w, h = 960, 540
    img = Image.new("RGB", (w, h), color=(245, 245, 245))
    draw = ImageDraw.Draw(img)
    border = (200, 200, 200)
    draw.rectangle([8, 8, w - 9, h - 9], outline=border, width=3)

    font_title = _load_font(28)
    font_sub = _load_font(20)
    font_small = _load_font(16)

    draw.text((w // 2, h // 2 - 50), title, fill=(80, 80, 80), font=font_title, anchor="mm")
    draw.text((w // 2, h // 2 + 10), "Screenshot placeholder", fill=(120, 120, 120), font=font_sub, anchor="mm")
    draw.text((w // 2, h // 2 + 55), hint, fill=(150, 150, 150), font=font_small, anchor="mm")

    img.save(out_path, format="PNG")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    out_dir = Path(__file__).resolve().parent
    out_docx = out_dir / "Security-Advisory-Patches-Portal-Visual-Guide.docx"

    placeholders = [
        ("fig_web_asgardeo.png", "Asgardeo sign-in", "Capture your organization login page"),
        ("fig_web_explorer.png", "Portal — folder listing", "Home / directory view after login"),
        ("fig_web_preview.png", "Preview & download", "File selected + preview panel + Download"),
        ("fig_azure_browse.png", "Azure File Share — browse", "Portal: file share folder & files"),
        ("fig_azure_upload.png", "Azure File Share — upload", "Upload / drag-drop or Upload folder"),
        ("fig_azcopy_terminal.png", "AzCopy — example run", "Terminal showing azcopy copy with progress"),
    ]

    with tempfile.TemporaryDirectory(prefix="sap_docx_") as tmp:
        tmp_path = Path(tmp)
        images: dict[str, Path] = {}
        for fname, title, hint in placeholders:
            p = tmp_path / fname
            make_placeholder_image(p, title, hint)
            images[fname] = p

        doc = Document()

        title = doc.add_heading("Security Advisory Patches Portal", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph("Visual operations guide (web app, Azure Files, AzCopy)")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(14)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = meta.add_run(f"Generated: {date.today().isoformat()}\nReplace gray placeholder images with real screenshots in Word (right-click picture → Change Picture).")
        mr.font.size = Pt(10)
        mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_page_break()

        doc.add_heading("1. Introduction", level=1)
        doc.add_paragraph(
            "This guide complements the technical README in the repository. It walks through three workflows: "
            "end users signing into the web application and browsing advisories; administrators managing content "
            "in Azure File Share from the Azure portal; and bulk data movement using AzCopy for large migrations."
        )

        # --- Web app ---
        doc.add_heading("2. Web application", level=1)

        doc.add_heading("2.1 Sign in with Asgardeo", level=2)
        doc.add_paragraph(
            "Open the portal URL provided by your organization. You are redirected to Asgardeo (or your hosted "
            "WSO2 Identity Server / Asgardeo tenant) to authenticate. After a successful login, you reach the file explorer."
        )
        bullets = doc.add_paragraph("Suggested capture:", style="List Bullet")
        bullets.add_run(
            " Full browser window showing the Asgardeo (or IdP) login screen, without exposing test passwords in the shot."
        )
        doc.add_paragraph().add_run().add_picture(str(images["fig_web_asgardeo.png"]), width=Inches(6.3))
        add_caption(doc, "Figure 1 — Asgardeo (or identity provider) login screen.")

        doc.add_heading("2.2 Home page and directory structure", level=2)
        doc.add_paragraph(
            "The main view lists folders and files from the Azure file share root (or the current folder). "
            "Users open folders to navigate; breadcrumbs show the path. Capture the listing that your customers will see day-to-day."
        )
        doc.add_paragraph().add_run().add_picture(str(images["fig_web_explorer.png"]), width=Inches(6.3))
        add_caption(doc, "Figure 2 — Portal home or subfolder listing.")

        doc.add_heading("2.3 File preview and download", level=2)
        doc.add_paragraph(
            "Selecting a file opens the preview panel (for PDFs and common image types). Use the Download control "
            "to save the file locally. Include both the file list and the preview pane if possible."
        )
        doc.add_paragraph().add_run().add_picture(str(images["fig_web_preview.png"]), width=Inches(6.3))
        add_caption(doc, "Figure 3 — Preview panel and download action.")

        doc.add_page_break()

        # --- Azure portal ---
        doc.add_heading("3. Azure File Share (Azure portal)", level=1)
        doc.add_paragraph(
            "Content is stored in a Microsoft Azure Storage account (Azure Files). Administrators upload or reorganize "
            "files here; the web portal reads the same share. Ensure folder names and permissions match your runbooks."
        )

        doc.add_heading("3.1 Navigate to the file share", level=2)
        doc.add_paragraph(
            "In the Azure portal: Storage accounts → your account → Data storage → File shares → select the share "
            "used by the Security Advisory Patches backend (see backend configuration for the share name)."
        )

        doc.add_heading("3.2 View folders and files", level=2)
        doc.add_paragraph(
            "Open directories and confirm files appear as expected. This view should mirror what portal users see after authentication."
        )
        doc.add_paragraph().add_run().add_picture(str(images["fig_azure_browse.png"]), width=Inches(6.3))
        add_caption(doc, "Figure 4 — Azure portal: file share directory listing.")

        doc.add_heading("3.3 Upload files or folders", level=2)
        doc.add_paragraph(
            "Use Upload, drag-and-drop, or folder upload options in the portal (exact controls vary slightly by portal version). "
            "For many small files or deep trees, prefer AzCopy (section 4)."
        )
        doc.add_paragraph().add_run().add_picture(str(images["fig_azure_upload.png"]), width=Inches(6.3))
        add_caption(doc, "Figure 5 — Azure portal: uploading into the file share.")

        doc.add_page_break()

        # --- AzCopy ---
        doc.add_heading("4. Bulk migration with AzCopy", level=1)
        doc.add_paragraph(
            "AzCopy v10 is the recommended tool for large migrations: high throughput, resumable transfers, and scriptable "
            "runs when uploading thousands of files or very large payloads to Azure Files."
        )

        doc.add_heading("4.1 Install and documentation", level=2)
        doc.add_paragraph(
            "Download and install AzCopy following Microsoft’s guide: "
            "https://learn.microsoft.com/en-us/azure/storage/common/storage-use-azcopy-v10"
        )
        doc.add_paragraph(
            "File-share–specific examples: "
            "https://learn.microsoft.com/en-us/azure/storage/common/storage-use-azcopy-files"
        )

        doc.add_heading("4.2 Authorization", level=2)
        doc.add_paragraph(
            "Typical approach: generate a SAS token on the storage account or file share with Read/Write/List (and Create) "
            "permissions as required, append it to the destination URL as ?sv=… . "
            "For directory-level operations, Microsoft also supports Microsoft Entra ID with AzCopy; see “Authorize AzCopy” "
            "in the get-started article."
        )

        doc.add_heading("4.3 Upload a directory recursively (common migration)", level=2)
        doc.add_paragraph(
            "This pattern uploads a local folder and all subfolders into a path inside the file share. "
            "Use single quotes around paths on macOS/Linux; in Windows cmd.exe use double quotes."
        )
        add_code_block(
            doc,
            "azcopy copy '<LOCAL_DIRECTORY>' "
            "'https://<STORAGE_ACCOUNT>.file.core.windows.net/<FILE_SHARE>/<OPTIONAL_TARGET_DIR>?<SAS_TOKEN>' "
            "--recursive",
        )
        doc.add_paragraph(
            "If <OPTIONAL_TARGET_DIR> does not exist, AzCopy can create it (see Microsoft Learn). "
            "Omit the directory segment to target the share root (still include the SAS on the URL)."
        )

        doc.add_heading("4.4 Optional flags (SMB / metadata)", level=2)
        doc.add_paragraph(
            "For Windows/SMB scenarios you may preserve ACLs and file properties (when supported):"
        )
        add_code_block(
            doc,
            "azcopy copy '<LOCAL_DIRECTORY>' "
            "'https://<STORAGE_ACCOUNT>.file.core.windows.net/<FILE_SHARE>/<PATH>?<SAS>' "
            "--recursive --preserve-permissions=true --preserve-info=true",
        )
        doc.add_paragraph(
            "For files larger than 256 MB where you want Content-MD5 set, add --put-md5 to the copy command (see Microsoft Learn)."
        )

        doc.add_heading("4.5 Verify after migration", level=2)
        doc.add_paragraph(
            "Compare file counts or spot-check in the Azure portal, then open the Security Advisory Patches web portal "
            "and confirm folders and previews behave as expected."
        )

        doc.add_paragraph().add_run().add_picture(str(images["fig_azcopy_terminal.png"]), width=Inches(6.3))
        add_caption(doc, "Figure 6 — AzCopy command line (progress / completion output).")

        doc.add_page_break()
        doc.add_heading("5. Document maintenance", level=1)
        doc.add_paragraph(
            "Regenerate this file by running: python3 build_visual_guide_docx.py "
            "(from the docs folder). "
            "Update screenshots when the UI, Asgardeo branding, or Azure portal layout changes."
        )

        doc.save(str(out_docx))

    print(f"Wrote {out_docx}")


if __name__ == "__main__":
    main()
